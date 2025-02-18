# -*- coding: utf-8 -*-

import pdfminer.high_level
import re
import openai
from tqdm import tqdm
# import nltk
# nltk.download('punkt')
# from nltk.tokenize import sent_tokenize
import ebooklib
from ebooklib import epub
import os
import tempfile
import shutil
import time
from bs4 import BeautifulSoup
import configparser
from datetime import datetime
from pytz import timezone
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from io import StringIO
import random
import json
import docx
import zipfile
from lxml import etree
from docx import Document
import mobi
import pandas as pd

def get_docx_title(docx_filename):
    with zipfile.ZipFile(docx_filename) as zf:
        core_properties = etree.fromstring(zf.read("docProps/core.xml"))

    ns = {"cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
          "dc": "http://purl.org/dc/elements/1.1/",
          "dcterms": "http://purl.org/dc/terms/",
          "dcmitype": "http://purl.org/dc/dcmitype/",
          "xsi": "http://www.w3.org/2001/XMLSchema-instance"}

    title_elements = core_properties.findall("dc:title", ns)
    if title_elements:
        return title_elements[0].text
    else:
        return "Unknown title"


def get_pdf_title(pdf_filename):
    try:
        with open(pdf_filename, 'rb') as file:
            parser = PDFParser(file)
            document = PDFDocument(parser)
            if 'Title' in document.info:
                return document.info['Title']
            else:
                text = pdfminer.high_level.extract_text(file)
                match = re.search(r'(?<=\n)([^\n]+)(?=\n)', text)
                if match:
                    return match.group(1)
                else:
                    return "Unknown title"
    except:
        return "Unknown title"


def get_mobi_title(mobi_filename):
    try:
        metadata = mobi.read_metadata(mobi_filename)
        title = metadata.get("Title", None)
    except:
        return "Unknown title"


def convert_mobi_to_text(mobi_filename):
    # Extract MOBI contents to a temporary directory
    with tempfile.TemporaryDirectory() as tempdir:
        tempdir, filepath = mobi.extract(mobi_filename)

        # Find the HTML file in the temporary directory
        for root, _, files in os.walk(tempdir):
            for file in files:
                if file.endswith(".html"):
                    html_file = os.path.join(root, file)
                    break
            else:
                continue
            break
        else:
            raise FileNotFoundError("HTML file not found in the extracted MOBI contents")

        # Parse the HTML file with BeautifulSoup to get the text
        with open(html_file, "r", encoding="utf-8") as f:
            soup = BeautifulSoup(f.read(), "html.parser")
            text = soup.get_text()

    return text


def get_epub_title(epub_filename):
    try:
        book = epub.read_epub(epub_filename)
        metadata = book.get_metadata('DC', {})
        if metadata:
            if 'title' in metadata:
                return metadata['title'][0]
        else:
            return "Unknown title"
    except:
        return "Unknown title"

    # 读取option文件


import chardet

with open('settings.cfg', 'rb') as f:
    content = f.read()
    encoding = chardet.detect(content)['encoding']

with open('settings.cfg', encoding=encoding) as f:
    config_text = f.read()
    config = configparser.ConfigParser()
    config.read_string(config_text)

# 获取openai_apikey和language
openai_apikey = os.getenv("OPENAI_API_KEY") # config.get('option', 'openai-apikey')
# language_name = config.get('option', 'target-language')
prompt = config.get('option', 'prompt')
bilingual_output = config.get('option', 'bilingual-output')
reformat = False # config.get('option', 'reformat') # RB added
language_code = config.get('option', 'langcode')
api_proxy=config.get('option', 'openai-proxy')
# Get startpage and endpage as integers with default values
startpage = config.getint('option', 'startpage', fallback=1)
endpage = config.getint('option', 'endpage', fallback=-1)
# 设置译名表文件路径
transliteration_list_file = config.get('option', 'transliteration-list')
# 译名表替换是否开启大小写匹配？
case_matching = config.get('option', 'case-matching')
# If want to print translations to console in realtime
print_translations = False

# 设置openai的API密钥
openai.api_key = openai_apikey

# 将openai的API密钥分割成数组
key_array = openai_apikey.split(',')

def random_api_key():
    return random.choice(key_array)

def create_chat_completion(prompt, text, model="gpt-4", **kwargs):
    openai.api_key = random_api_key()
    return openai.ChatCompletion.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": f"{prompt}: \n{text}",
            }
        ],
        **kwargs
    )

import argparse

# 如果配置文件有写，就设置api代理
if len(api_proxy) == 0:
    print("未检测到OpenAI API 代理，当前使用api地址为: " + openai.api_base)
else:
    api_proxy_url = api_proxy + "/v1"
    openai.api_base = os.environ.get("OPENAI_API_BASE", api_proxy_url)
    print("正在使用OpenAI API 代理，代理地址为: "+openai.api_base)

# 创建参数解析器
parser = argparse.ArgumentParser()
parser.add_argument("filename", help="Name of the input file")
parser.add_argument("--test", help="Only translate the first 3 short texts", action="store_true")
# 是否使用译名表？
parser.add_argument("--tlist", help="Use the translated name table", action="store_true")
args = parser.parse_args()

# 获取命令行参数
filename = args.filename
base_filename, file_extension = os.path.splitext(filename)
new_filename = base_filename + "_translated.epub"
new_filenametxt = base_filename + "_translated.txt"
jsonfile = base_filename + "_process.json"
# 从文件中加载已经翻译的文本
translated_dict = {}
try:
    with open(jsonfile, "r", encoding="utf-8") as f:
        print("Loading translated text from " + jsonfile)
        translated_dict = json.load(f)
except FileNotFoundError:
    print("File Not found")
    pass


def convert_docx_to_text(docx_filename):
    doc = docx.Document(docx_filename)

    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"

    return text


def convert_epub_to_text(epub_filename):
    # 打开epub文件
    book = epub.read_epub(epub_filename)

    # 获取所有文本
    text = ""
    for item in book.get_items():
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            # 使用BeautifulSoup提取纯文本
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text += re.sub(r'\n+', '\n', soup.get_text().strip())

    # 返回文本
    return text

def convert_EPUB_to_DOC(self, epub_filename, outfile):
        path_infile = self.dataDir + infile

        options = aw.loading.EpubLoadOptions()

        # Open EPUB document

        document = Document(path_infile, options)

        option = DocSaveOptions()
        option.Format = DocSaveOptions.DocFormat.DocX

        # Save the file into MS Word document format

        document.Save(path_outfile, option)
        print(infile + " converted into " + outfile)

# RB replaces text_to_epub to preserve original epub formatting
def merge_to_new_epub(original_book, translated_dict, new_filename, language_code='en', title="Title"):
    # Create a new epub book
    new_book = epub.EpubBook()

    # Step 1: Extract original metadata and contents
    # Metadata
    new_book.set_identifier(original_book.get_metadata('DC', 'identifier')[0][0])
    new_book.set_title(title)
    new_book.set_language(original_book.language)

    cover_page = ebooklib.epub.EpubCoverHtml(title='Cover', file_name='cover.xhtml')
    cover_page.content = '<your HTML content here>'

    # Cover Image (if available)
    # cover_item = original_book.get_cover_item()
    # if cover_item:
    #     new_book.set_cover("image.jpg", cover_item.content)
    #     print("Applied cover image")
    
    # Initialize new table of contents and spine
    new_toc = []
    new_spine = ['nav']
    
    # Get original table of contents (TOC)
    # original_toc = original_book.get_items_of_type(ebooklib.ITEM_NAVIGATION) # original_toc = original_book.toc

    # No need for new_toc or new_spine as separate lists
    # Simply copy over the original TOC and spine
    new_book.toc = original_book.toc
    new_book.spine = original_book.spine

    # Loop through all the items in the original book
    for item in original_book.get_items():
        # Check if the item is a document (i.e., a chapter or section)
        if item.get_type() == ebooklib.ITEM_DOCUMENT:
            chapter_content = item.content.decode('utf-8').strip()  # Decoding and stripping leading/trailing white spaces

            chapter_title = item.title if item.title else "Unknown Title"
            print(f"Found chapter: {chapter_title}")

            # Create a new chapter with translated content, falling back to original if not found
            translated_content = translated_dict.get(chapter_content, chapter_content)

            new_chapter = ebooklib.epub.EpubHtml(title=chapter_title, content=translated_content, lang=language_code)
            new_chapter.content = translated_content.encode('utf-8')  # Make sure to encode back to bytes

            # Add new chapter to the book
            new_book.add_item(new_chapter)


    # Write out new ePub file
    ebooklib.epub.write_epub(new_filename, new_book)

# RB wrote to try fix formatting issues
def text_to_epub(original_book, text, filename, language_code='en', title="Title"):
    print("Writing translated text to epub \n" + filename)

    # text = text.replace("\n", "<br>") # USING CSS NAO

    # 创建epub书籍对象
    new_book = epub.EpubBook()

    # 设置元数据
    new_book.set_identifier(str(random.randint(100000, 999999)))
    new_book.set_title(title)
    new_book.set_language(language_code)

    # Create CSS style for paragraph spacing
    css_style = '''
    @namespace epub "http://www.idpf.org/2007/ops";
    body {
        font-family: Arial, sans-serif;
    }
    p {
        text-align: justify;
        margin-bottom: 1.5em;
    }
    '''
    default_css = epub.EpubItem(uid="style_default", file_name="style/default.css", media_type="text/css", content=css_style)
    # apply css styling across the content of the new_book
    new_book.add_item(default_css)

    # Create chapter
    chapter1 = epub.EpubHtml(title='Chapter 1', file_name='chap_1.xhtml', lang=language_code)
    chapter1.content = '<html><head><link rel="stylesheet" type="text/css" href="style/default.css" /></head><body>' + text + '</body></html>'

    # 将章节添加到书籍中
    new_book.add_item(chapter1)

    # 添加toc
    new_book.toc = (epub.Link('chap_1.xhtml', 'Chapter 1', 'chap_1'),)
    # 设置书脊顺序
    new_book.spine = ['nav', chapter1]

    # 添加导航
    new_book.add_item(epub.EpubNcx())
    new_book.add_item(epub.EpubNav())


    # COVER COPY 
    # OG: new_book.set_cover('image.jpg', open('image.jpg', 'rb').read())
    # Get the cover item from the original book
    # Look for the cover item in the original book
    cover_item = None
    for item in original_book.get_items():
        if item.get_type() == ebooklib.ITEM_COVER:
            cover_item = item
            break

    # If a cover item is found, copy it to the new book
    if cover_item:
        # Create a new EpubCover item for the new book
        new_cover_item = ebooklib.epub.EpubCover(
            uid='cover-img',
            file_name=cover_item.file_name  # you can give a new name if you like
        )
        new_cover_item.content = cover_item.content

        # Add the new cover item to the new book
        new_book.add_item(new_cover_item)


    # Final step: write out the new book
    epub.write_epub(filename, new_book, {})


# 将PDF文件转换为文本
# For PDF files
def get_total_pages(pdf_filename):
    with open(pdf_filename, 'rb') as file:
        parser = PDFParser(file)
        document = PDFDocument(parser)
        return len(list(PDFPage.create_pages(document)))


def convert_pdf_to_text(pdf_filename, start_page=1, end_page=-1):
    if end_page == -1:
        end_page = get_total_pages(pdf_filename)
        # print("Total pages of the file:"+ str(end_page))
        # print("Converting PDF from:"+ str(start_page)+" to "+ str(end_page) + " page")
        text = pdfminer.high_level.extract_text(pdf_filename, page_numbers=list(range(start_page - 1, end_page)))
    else:
        # print("Converting PDF from:"+ str(start_page)+" to "+ str(end_page) + " page")
        text = pdfminer.high_level.extract_text(pdf_filename, page_numbers=list(range(start_page - 1, end_page)))
    return text


# 将文本分成不大于1024字符的短文本list
def split_text(text):
    sentence_list = re.findall(r'.+?[。！？!?.]', text)

    # 初始化短文本列表
    short_text_list = []
    # 初始化当前短文本
    short_text = ""
    # 遍历句子列表
    for s in sentence_list:
        # 如果当前短文本加上新的句子长度不大于1024，则将新的句子加入当前短文本
        if len(short_text + s) <= 1024:
            short_text += s
        # 如果当前短文本加上新的句子长度大于1024，则将当前短文本加入短文本列表，并重置当前短文本为新的句子
        else:
            short_text_list.append(short_text)
            short_text = s
    # 将最后的短文本加入短文本列表
    short_text_list.append(short_text)
    return short_text_list


# 将句号替换为句号+回车
def return_text(text):
    # text = text.replace(". ", ".\n")
    text = text.replace("。", "。\n")
    text = text.replace("！", "！\n")
    return text


# Initialize a count variable of tokens cost.
cost_tokens = 0


# 翻译短文本
def translate_text(text):
    global cost_tokens

    # 调用openai的API进行翻译
    try:
        completion = create_chat_completion(prompt, text)
        t_text = (
            completion["choices"][0]
            .get("message")
            .get("content")
            .encode("utf8")
            .decode()
        )
        # Get the token usage from the API response
        cost_tokens += completion["usage"]["total_tokens"]

    except Exception as e:
        import time
        # TIME LIMIT for open api please pay
        sleep_time = 60
        time.sleep(sleep_time)
        print(e, f"will sleep  {sleep_time} seconds")

        completion = create_chat_completion(prompt, text)
        t_text = (
            completion["choices"][0]
            .get("message")
            .get("content")
            .encode("utf8")
            .decode()
        )
        # Get the token usage from the API response
        cost_tokens += completion["usage"]["total_tokens"]

    return t_text


def translate_and_store(text):
    # print("Translating text: " + text)
    # If the text has already been translated, directly return the translation result
    if text in translated_dict:
        return translated_dict[text]
    
    ##########################################################
    # TRANSLATING THE TEXT HERE
    ##########################################################
    translated_text = translate_text(text)
    translated_dict[text] = translated_text

    # Save the dictionary as a JSON file
    with open(jsonfile, "w", encoding="utf-8") as f:
        json.dump(translated_dict, f, ensure_ascii=False, indent=4)

    return translated_text


def text_replace(long_string, xlsx_path, case_sensitive):
    # 读取excel文件，将第一列和第二列分别存为两个列表
    df = pd.read_excel(xlsx_path)
    old_words = df.iloc[:, 0].tolist()
    new_words = df.iloc[:, 1].tolist()
    # 对旧词列表按照长度降序排序，并同步调整新词列表
    old_words, new_words = zip(*sorted(zip(old_words, new_words), key=lambda x: len(x[0]), reverse=True))
    # 遍历两个列表，对字符串进行替换
    for i in range(len(old_words)):
        # 如果不区分大小写，就将字符串和被替换词都转为小写
        if not case_sensitive:
            lower_string = long_string.lower()
            lower_old_word = old_words[i].lower()
            # 使用正则表达式进行替换，注意要保留原字符串的大小写
            long_string = re.sub(r"\b" + lower_old_word + r"\b", new_words[i], long_string, flags=re.IGNORECASE)
        # 如果区分大小写，就直接使用正则表达式进行替换
        else:
            long_string = re.sub(r"\b" + old_words[i] + r"\b", new_words[i], long_string)
    # 返回替换后的字符串
    return long_string


text = ""

# 根据文件类型调用相应的函数
if filename.endswith('.pdf'):
    print("Converting PDF to text")
    title = get_pdf_title(filename)
    with tqdm(total=10, desc="Converting PDF to text") as pbar:
        for i in range(10):
            text = convert_pdf_to_text(filename, startpage, endpage)
            pbar.update(1)

elif filename.endswith('.epub'):
    print("Converting epub to text")
    book = epub.read_epub(filename)

elif filename.endswith('.txt'):

    with open(filename, 'r', encoding='utf-8') as file:
        text = file.read()

    title = os.path.basename(filename)

elif filename.endswith('.docx'):
    print("Converting DOCX file to text")
    title = get_docx_title(filename)
    with tqdm(total=10, desc="Converting DOCX to text") as pbar:
        for i in range(10):
            text = convert_docx_to_text(filename)
            pbar.update(1)

elif filename.endswith('.mobi'):
    print("Converting MOBI file to text")
    title = get_mobi_title(filename)
    with tqdm(total=10, desc="Converting MOBI to text") as pbar:
        for i in range(10):
            text = convert_mobi_to_text(filename)
            pbar.update(1)
else:
    print("Unsupported file type")

# RB added
if reformat: 
    print(">>> Reformatting text <<<")

    with tqdm(total=4, desc="Processing and Writing EPUB") as pbar:
        # Step 1: Read the JSON file
        with open("book_process.json", 'r', encoding='utf-8') as f:
            translated_dict = json.load(f)
        pbar.update(1)

        # Step 2: Assume the translated text is the values in the dictionary
        # Split each value by \n and wrap each resulting string in <p> tags
        reformatted_text = ""
        for paragraph in translated_dict.values():
            lines = paragraph.split('\n')
            reformatted_text += "".join([f"<p>{line}</p>" for line in lines])
        pbar.update(1)

        # Step 3: Create EPUB
        book = epub.EpubBook()
        book.set_identifier(str(random.randint(100000, 999999)))
        book.set_title("PeregrinosEncuentros en el Camino de Santiago")
        book.set_language(language_code)

        c = epub.EpubHtml(title='Translated Chapter', file_name='chap_translated.xhtml', lang=language_code)
        c.content = reformatted_text
        book.add_item(c)

        # Append the new chapter to the existing TOC and spine
        book.toc.append(epub.Link('chap_translated.xhtml', 'Translated Chapter', 'chap_translated'))
        book.spine.append(c)
        
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        epub.write_epub(new_filename, book, {})
        pbar.update(1)

# Gets data as a "book" for epub, or as string format for other file types
if filename.endswith('.pdf'):
    title = get_pdf_title(filename)
    with tqdm(total=10, desc="Converting PDF to text") as pbar:
        for i in range(10):
            text = convert_pdf_to_text(filename, startpage, endpage)
            pbar.update(1)

elif filename.endswith('.epub'):
    book = epub.read_epub(filename)
    title = get_epub_title(filename)
    print ("Epub Title: " + title)
    print("\nConverted epub to text as a 'book'")
    input("Press Enter to continue...")


elif filename.endswith('.txt'):
    with open(filename, 'r', encoding='utf-8') as file:
        text = file.read()
    title = os.path.basename(filename)

elif filename.endswith('.docx'):
    print("Converting DOCX file to text")
    title = get_docx_title(filename)
    with tqdm(total=10, desc="Converting DOCX to text") as pbar:
        for i in range(10):
            text = convert_docx_to_text(filename)
            pbar.update(1)

elif filename.endswith('.mobi'):
    print("Converting MOBI file to text")
    title = get_mobi_title(filename)
    with tqdm(total=10, desc="Converting MOBI to text") as pbar:
        for i in range(10):
            text = convert_mobi_to_text(filename)
            pbar.update(1)

else:
    print("Unsupported file type")

img_html_added = False

import re

# Peform the translation
if filename.endswith('.epub'):
    print("\nTranslating epub book")
    #print First translated sentence from the json file
    if translated_dict.__len__() > 0:
        print("\t(First translated sentence) " + translated_dict[list(translated_dict.keys())[0]])
        print("\t(Last translated sentence) " + translated_dict[list(translated_dict.keys())[-1]])
    input("Press Enter to continue...")

    # 获取所有章节
    items = book.get_items()

    # 遍历所有章节
    translated_all = ''
    count = 0

    for item in tqdm(items):
        if item.get_type() == ebooklib.ITEM_DOCUMENT:

            # Data structure for parsed html or xml
            soup = BeautifulSoup(item.get_content(), 'html.parser')
            text = soup.get_text().strip()
            img_html = ''
            img_tags = soup.find_all('img')

             # html tags for images in this chapter
            for img_tag in img_tags:
                img_html += str(img_tag) + '<br>'
                
            if not text: # ... its not text dumb dumb
                continue

            # remove extra spaces
            text = re.sub(r"\s+", " ", text)

            # transliterate the text if it is not in the translated_dict
            if args.tlist:
                text = text_replace(text, transliteration_list_file, case_matching)

            # split the text into a "short text" list
            short_text_list = split_text(text)

            # RB save the short_text_list to a new raw tex file
            with open("raw_text.txt", "w", encoding="utf-8") as f:
                f.write(text)
                
            if args.test:
                short_text_list = short_text_list[:3]

            translated_text = ""

            # unformatted translated text with all tags in the string
            unformatted_translated_text = ""

            # TRAVERSE THE LIST OF SHORT TEXTS, TRANSLATING EACH SHORT TEXT IN TURN
            for short_text in tqdm(short_text_list):
                count += 1

                translated_short_text = translate_and_store(short_text)
                short_text = return_text(short_text)
                unformatted_translated_text = translated_short_text
                translated_short_text = return_text(translated_short_text)

                if bilingual_output.lower() == 'true':
                    translated_text += f"{short_text}<br>\n{translated_short_text}<br>\n"
                else:
                    translated_text += f"{translated_short_text}<br>\n"
            
            item.set_content((img_html + translated_text.replace('\n', '<br>')).encode('utf-8'))

            # print translation to screen
            print("Translated segment!")
            print("\033[95m\nSpanish: " + short_text + "..." + "033[0m")
            print("\033[93m\mEnglish: " + translated_short_text + "..." + "\n\033[0m")

            # add the translated text to the overall translated text
            translated_all += translated_text
            if args.test and count >= 3:
                break

    # SAVE TO FILE!
    epub.write_epub(new_filename, book, {})
    with open(new_filenametxt, "w", encoding="utf-8") as f:
        f.write(translated_all)

else:
    print("\n(NOT epub) Processing escape characters, removing exra spaces, performing transliterations/lang replacements, splitting text")
    input("Press Enter to continue...")

    # 将所有回车替换为空格
    text = text.replace("\n", " ")

    # 将多个空格替换为一个空格
    import re

    text = re.sub(r"\s+", " ", text)

    # 如果设置了译名表替换，则对文本进行翻译前的替换
    if args.tlist:
        text = text_replace(text, transliteration_list_file, case_matching)

    # 将文本分成不大于1024字符的短文本list
    short_text_list = split_text(text)
    # dump the short_text_list to a new raw tex file
    with open("raw_text.txt", "w", encoding="utf-8") as f:
        f.write(text)
   
    if args.test:
        short_text_list = short_text_list[:3]
    # 初始化翻译后的文本
    translated_text = ""

    # Traverse the list of short texts, translating each short text in turn
    for short_text in tqdm(short_text_list):
        print(return_text(short_text))
        # Translate the current short text
        translated_short_text = translate_and_store(short_text)
        short_text = return_text(short_text)
        translated_short_text = return_text(translated_short_text)
        # Add the current short text and its translation to the overall text
        if bilingual_output.lower() == 'true':
            translated_text += f"{short_text}\n{translated_short_text}\n"
        else:
            translated_text += f"{translated_short_text}\n"
        # print(short_text)
        print(translated_short_text)

def generate_line(length):
    return ''.join(random.choice(characters) for _ in range(length))
def funky_finalizer(duration = 3):
    import random

    # Characters to use
    characters = "abcdefghijklmnopqrstuvwxyzABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789!@#$%^&*()-=_+[]{}|;':,.<>/?`~"

    lines = [generate_line(line_length) for _ in range(num_lines)]
    start_time = time.time()
    while True:
        current_time = time.time()
        if current_time - start_time >= duration:
            break
        # Print the lines
        for line in lines:
            print(line)
        
        # Update a random line to create the scrolling effect
        lines[random.randint(0, num_lines-1)] = generate_line(line_length)
        
        # Delay
        time.sleep(0.1)
        
        # Clear the console screen
        print("\033c", end="")

    # Display the cryptic message
    print(">> Data stream ciphered and archived to matrix_data.epub")
    print(">> WARNING: Disclosure of this data will rupture the digital mainframe, proceed with extreme caution.")

    if __name__ == "__main__":
        matrix_effect()


funky_finalizer()

# Write the translated text to a TXT file as well, in case there are issues with the EPUB plugin
with open(new_filenametxt, "w", encoding="utf-8") as f:
    f.write(translated_text)

cost = cost_tokens / 1000 * 0.002
print(f"Translation completed. Total cost: {cost_tokens} tokens, ${cost}.")

# Get human-readable time string in EST for saving the backup files
eastern = timezone('US/Eastern')
utc_time = datetime.now(timezone('UTC'))
est_time = utc_time.astimezone(eastern)
time_string_est = est_time.strftime("-%Y-%m-%d-%H-%M-%S")

# Backup the JSON file
try:
    # copy the json and append human readable date/time in current timzeone to the filename
    copy = shutil.copy(jsonfile, jsonfile + time_string_est)
    # move the duplicate json to folder
    shutil.move(copy, "Processed Jsons")
    # print the success message
    print(f"File '{jsonfile}' has been backed up to the 'Processed Jsons' folder.")
except FileNotFoundError:
    print(f"File '{jsonfile}' not found, or backup folder 'Processed Jsons' does not exist. Backup failed.")

# Backup the epub file
try:
    # copy the epub and append human readable date/time in current timzeone to the filename
    copy = shutil.copy(new_filename, new_filename + time_string_est)
    # move the duplicate epub to folder
    shutil.move(copy, "Translated Epubs")
    # print the success message
    print(f"File '{new_filename}' has been backed up to the 'Translated Epubs' folder.")
except FileNotFoundError:
    print(f"File '{new_filename}' not found, or backup folder 'Translated Epubs' does not exist. Backup failed.")

print("\nAll done! Enjoy your translated book!")